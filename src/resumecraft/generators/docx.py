"""DOCX resume generator, built on python-docx.

Note on borders: python-docx's `CT_PPr` does NOT expose a
`get_or_add_pBdr()` convenience method (unlike `get_or_add_pStyle` or
`get_or_add_numPr`, which do exist). Paragraph borders simply aren't
covered by the high-level API, so the only way to draw a rule under a
paragraph is to build the `<w:pBdr><w:bottom .../></w:pBdr>` XML by hand
and attach it to the paragraph's `<w:pPr>`. That is centralized in
`_add_bottom_border` below and used everywhere a divider or section
underline is needed, instead of being re-implemented (incorrectly) at
each call site.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from resumecraft.exceptions import GenerationError
from resumecraft.generators.base import BaseGenerator
from resumecraft.utils import rgb_tuple

logger = logging.getLogger("resumecraft")


def _primary_font(font_family: str) -> str:
    """DOCX needs a single font name; take the first one from a CSS-style stack."""
    first = font_family.split(",")[0].strip()
    return first.strip("'\"") or "Calibri"


class DOCXGenerator(BaseGenerator):
    """Generates a native .docx resume."""

    def __init__(
        self,
        data,
        *,
        theme_name: str = "clarity",
        lang: str = "ru",
        accent_color: str = "2B6CB0",
        font_family: str = "Calibri, 'Segoe UI', Arial, sans-serif",
    ) -> None:
        super().__init__(
            data,
            theme_name=theme_name,
            lang=lang,
            accent_color=accent_color,
            font_family=font_family,
        )
        self.font_name = _primary_font(font_family)
        self.accent_rgb = RGBColor(*rgb_tuple(self.accent_color))

    # ---- low-level helpers -------------------------------------------------

    def _add_bottom_border(self, paragraph, size: str = "6") -> None:
        """Add (or extend) a bottom border on a paragraph. See module docstring."""
        pPr = paragraph._p.get_or_add_pPr()

        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = OxmlElement("w:pBdr")
            pPr.append(pBdr)

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), size)
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), self.accent_color)
        pBdr.append(bottom)

    def _add_horizontal_line(self, doc: Document) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        self._add_bottom_border(p, size="6")

    def _add_section_title(self, doc: Document, text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = self.accent_rgb
        run.font.name = self.font_name
        self._add_bottom_border(p, size="4")
        return p

    # ---- document setup ------------------------------------------------------

    def _setup_document(self, doc: Document) -> None:
        style = doc.styles["Normal"]
        style.font.name = self.font_name
        style.font.size = Pt(10.5)
        for section in doc.sections:
            section.top_margin = Pt(36)
            section.bottom_margin = Pt(36)
            section.left_margin = Pt(48)
            section.right_margin = Pt(48)

    def _add_header(self, doc: Document) -> None:
        name_p = doc.add_paragraph()
        name_p.paragraph_format.space_after = Pt(0)
        run = name_p.add_run(self.data.full_name)
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = self.font_name

        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_after = Pt(6)
        run = title_p.add_run(self.data.title)
        run.italic = True
        run.font.size = Pt(12.5)
        run.font.color.rgb = self.accent_rgb
        run.font.name = self.font_name

        contact_bits = [self.data.location, self.data.phone, self.data.email]
        for extra in (self.data.linkedin, self.data.github, self.data.website):
            if extra:
                contact_bits.append(extra)
        if self.data.salary:
            contact_bits.append(f"Зарплата: {self.data.salary}")

        contact_line = "  |  ".join(b for b in contact_bits if b)
        contact_p = doc.add_paragraph()
        contact_p.paragraph_format.space_after = Pt(8)
        run = contact_p.add_run(contact_line)
        run.font.size = Pt(9.5)
        run.font.name = self.font_name

        self._add_horizontal_line(doc)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def _add_summary(self, doc: Document) -> None:
        self._add_section_title(doc, self.labels["summary"])
        p = doc.add_paragraph(self.data.summary)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def _add_experience(self, doc: Document) -> None:
        if not self.data.experience:
            return
        self._add_section_title(doc, self.labels["experience"])
        for exp in self.data.experience:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_after = Pt(0)
            run = title_p.add_run(f"{exp.position} — {exp.company}")
            run.bold = True
            run.font.name = self.font_name

            meta_bits = [exp.duration_text(self.present_label())]
            if exp.location:
                meta_bits.append(exp.location)
            meta_p = doc.add_paragraph(" · ".join(meta_bits))
            meta_p.paragraph_format.space_after = Pt(3)
            for run in meta_p.runs:
                run.italic = True
                run.font.size = Pt(9.5)
                run.font.name = self.font_name

            for line in exp.description:
                bullet_p = doc.add_paragraph(line, style="List Bullet")
                bullet_p.paragraph_format.space_after = Pt(2)

            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def _add_education(self, doc: Document) -> None:
        if not self.data.education:
            return
        self._add_section_title(doc, self.labels["education"])
        for edu in self.data.education:
            degree = edu.degree if not edu.field else f"{edu.degree}, {edu.field}"
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_after = Pt(0)
            run = title_p.add_run(degree)
            run.bold = True
            run.font.name = self.font_name

            meta_bits = [edu.institution, edu.duration_text(self.present_label())]
            if edu.gpa:
                meta_bits.append(f"GPA {edu.gpa}")
            meta_p = doc.add_paragraph(" · ".join(meta_bits))
            meta_p.paragraph_format.space_after = Pt(6)
            for run in meta_p.runs:
                run.italic = True
                run.font.size = Pt(9.5)
                run.font.name = self.font_name

    def _add_skills(self, doc: Document) -> None:
        if not self.data.skills:
            return
        self._add_section_title(doc, self.labels["skills"])
        for cat in self.data.skills:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(f"{cat.category}: ")
            run.bold = True
            run.font.name = self.font_name
            run2 = p.add_run(", ".join(cat.items))
            run2.font.name = self.font_name

    def _add_projects(self, doc: Document) -> None:
        if not self.data.projects:
            return
        self._add_section_title(doc, self.labels["projects"])
        for proj in self.data.projects:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_after = Pt(0)
            run = title_p.add_run(proj.name)
            run.bold = True
            run.font.name = self.font_name
            if proj.url:
                title_p.add_run(f"  ({proj.url})").font.size = Pt(9)

            if proj.description:
                doc.add_paragraph(proj.description).paragraph_format.space_after = Pt(2)
            if proj.technologies:
                tech_p = doc.add_paragraph(", ".join(proj.technologies))
                tech_p.paragraph_format.space_after = Pt(6)
                for run in tech_p.runs:
                    run.italic = True
                    run.font.size = Pt(9.5)

    def _add_languages(self, doc: Document) -> None:
        if not self.data.languages:
            return
        self._add_section_title(doc, self.labels["languages"])
        text = ", ".join(
            f"{l.language} ({l.level})" if l.level else l.language for l in self.data.languages
        )
        doc.add_paragraph(text)

    def _add_certifications(self, doc: Document) -> None:
        if not self.data.certifications:
            return
        self._add_section_title(doc, self.labels["certifications"])
        for cert in self.data.certifications:
            doc.add_paragraph(cert, style="List Bullet")

    def _add_additional(self, doc: Document) -> None:
        if not self.data.additional:
            return
        self._add_section_title(doc, self.labels["additional"])
        doc.add_paragraph(self.data.additional)

    # ---- entry point ---------------------------------------------------------

    def _generate(self, output_path: Path) -> Path:
        try:
            doc = Document()
            self._setup_document(doc)
            self._add_header(doc)
            self._add_summary(doc)
            self._add_experience(doc)
            self._add_education(doc)
            self._add_skills(doc)
            self._add_projects(doc)
            self._add_languages(doc)
            self._add_certifications(doc)
            self._add_additional(doc)
            doc.save(str(output_path))
        except GenerationError:
            raise
        except Exception as e:
            raise GenerationError(f"Failed to generate DOCX resume: {e}") from e

        logger.info("DOCX saved: %s", output_path)
        return output_path
